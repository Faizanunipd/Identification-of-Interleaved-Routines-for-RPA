import os
import random
import numpy as np
import pandas as pd
import scipy.stats as stats
from tabulate import tabulate
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from collections import defaultdict
from IPython.display import display, HTML
from matplotlib.ticker import MaxNLocator
import re

import pm4py
from pm4py.algo.conformance.tokenreplay import algorithm as token_replay
from pm4py.objects.petri_net.importer.variants import pnml as pnml_importer


######################################################################################################################
####                                                   Main Functions                                             ####
######################################################################################################################

# --- Main process function ---
def process_random_logs(logs, document, target_datetime_str="2024-06-07 21:34:25.493914+00:00", variance_criteria='max', specified_routines=None):
    while True:
        # Step 1: Select routines
        if specified_routines is not None:
            selected_routines = sorted(specified_routines)
        else:
            selected_routines = select_random_routines(logs)
        
        # Step 2: Process logs
        # Capture original trace counts BEFORE any processing
        original_logs = [logs[logs['routine_type'] == rt].copy() for rt in selected_routines]
        traces_per_routine_original = {}
        for i, log in enumerate(original_logs):
            traces_per_routine_original[selected_routines[i]] = log['case:concept:name'].nunique()
        
        selected_logs = select_and_renumber_cases_all(logs, selected_routines)
        adjusted_logs = adjust_timestamps_to_target(selected_logs, target_datetime_str)
        adjusted_logs = force_first_trace_to_target(adjusted_logs, target_datetime_str)
        time_diffs, variances = calculate_inter_trace_gaps(adjusted_logs)
        reference_log_idx, best_dist, best_params, sorted_time_diffs = fit_reference_distribution(time_diffs, variances, adjusted_logs, variance_criteria)
        updated_logs = shift_traces_by_distribution(adjusted_logs, reference_log_idx, best_dist, best_params, sorted_time_diffs)
        final_log = pd.concat(updated_logs, ignore_index=True).sort_values(by=["time:timestamp"]).reset_index(drop=True)
        
        # Trace counts
        num_traces_before = final_log['case:concept:name'].nunique()
        traces_per_routine_before = final_log.groupby('routine_type')['case:concept:name'].nunique()
        num_traces_after = num_traces_before
        traces_per_routine_after = traces_per_routine_before
        
        # Compute interleaving
        interleave_counts = count_trace_interleaving_cases(final_log)

        # Check acceptance conditions
        if (
            num_traces_after >= max(1, int(0.05 * num_traces_before))
        ):
            document = display_and_store_selected_routines(selected_routines, logs, document)
            
            # Collect metadata for iteration summary
            iteration_metadata = {
                'selected_routines': selected_routines,
                'traces_before_renumbering': traces_per_routine_original,
                'traces_after_renumbering': traces_per_routine_after.to_dict(),
                'interleaving_counts': interleave_counts
            }
            
            # Sufficient traces, proceed to output
            final_log = final_log[['concept:name', 'time:timestamp', 'case:concept:name', 'routine_type', 'log_number']]
            segment_log = final_log.copy()
            unsegment_log = final_log.copy()
            unsegment_log['case:concept:name'] = 1
            
            return segment_log, unsegment_log, selected_routines, document, iteration_metadata
        else:
            if specified_routines is not None:
                return None, None, None, document, None
            continue


######################################################################################################################
####                                            Functions Called by process_random_logs                           ####
######################################################################################################################

def select_random_routines(logs):
    """Select random routines based on trace count and activity criteria"""
    # Static variable to track if this is the first call
    if not hasattr(select_random_routines, 'first_call'):
        select_random_routines.first_call = True
    
    # Calculate trace counts and unique activities for each routine type
    trace_counts = logs.groupby('routine_type')['case:concept:name'].nunique()
    unique_activities = logs.groupby('routine_type')['concept:name'].nunique()
    
    # Only print detailed info on first call
    if select_random_routines.first_call:
        print("*"*50)
        all_trace_couts = sum(logs.groupby('routine_type')['case:concept:name'].nunique())
        print(f"Total routine types before filtering {len(trace_counts)} with trace count {all_trace_couts}")
        
        # Filter by trace count
        trace_mean = int(sum(trace_counts) / len(trace_counts))
        # trace_median = int(trace_counts.median())
        # trace_first_quartile = int(np.percentile(trace_counts, 75))
        trace_eligible = trace_counts[trace_counts > trace_mean].index.tolist()
        print(f"Trace mean: {trace_mean:.1f}")
        print(f"Routine types after trace count filter (> {trace_mean:.1f}): {len(trace_eligible)}")
        
        # Filter by unique activities
        activity_median = unique_activities.median()
        activity_eligible = unique_activities[unique_activities > activity_median].index.tolist()
        print(f"Routine types after activity count filter (> {activity_median:.1f}): {len(activity_eligible)}")
        
        # Get intersection of both filters
        eligible_routines = list(set(trace_eligible) & set(activity_eligible))
        print(eligible_routines)
        eligible_trace_couts = sum(logs[logs['routine_type'].isin(eligible_routines)].groupby('routine_type')['case:concept:name'].nunique())
        print(f"Routine types after both filters are {len(eligible_routines)} with {eligible_trace_couts} traces")
        # print(f"Routine types after both filters: {len(eligible_routines)}")
        
        select_random_routines.first_call = False
        print("*"*50)
    else:
        # For subsequent calls, just calculate the filters without printing
        trace_mean = int(sum(trace_counts) / len(trace_counts))
        trace_eligible = trace_counts[trace_counts > trace_mean].index.tolist()
        activity_median = unique_activities.median()
        activity_eligible = unique_activities[unique_activities > activity_median].index.tolist()
        eligible_routines = list(set(trace_eligible) & set(activity_eligible))
    
    if len(eligible_routines) < 3:
        if select_random_routines.first_call:
            print(f"Warning: Only {len(eligible_routines)} routines meet criteria, selecting all available")
        selected_routines = sorted(eligible_routines)
    else:
        selected_routines = sorted(random.sample(eligible_routines, 3))
    
    print(f"\nSelected routines: {selected_routines}")
    return selected_routines


def select_and_renumber_cases_all(logs, selected_routines):
    """Keep all traces from each routine and renumber case IDs sequentially"""
    selected_logs = [logs[logs['routine_type'] == rt].copy() for rt in selected_routines]
    
    print(f"Original trace counts per routine:")
    for i, log in enumerate(selected_logs):
        print(f"  Routine {selected_routines[i]}: {log['case:concept:name'].nunique()} traces")
    
    next_case_id = 1
    for i, log in enumerate(selected_logs):
        # Get all unique case IDs for this routine (no limiting)
        unique_cases = log['case:concept:name'].unique()
        
        # Create mapping to new sequential integer IDs
        case_mapping = {old: new for old, new in zip(unique_cases, range(next_case_id, next_case_id + len(unique_cases)))}
        log['case:concept:name'] = log['case:concept:name'].map(case_mapping)
        selected_logs[i] = log
        next_case_id += len(unique_cases)
        
        print(f"  Routine {selected_routines[i]} after renumbering: {len(unique_cases)} traces (case IDs: {min(case_mapping.values())}-{max(case_mapping.values())})")
    
    return selected_logs


def select_and_renumber_cases_minimum(logs, selected_routines):
    """Limit each routine to minimum traces and renumber case IDs sequentially"""
    selected_logs = [logs[logs['routine_type'] == rt].copy() for rt in selected_routines]
    
    print(f"\n=== BEFORE select_and_renumber_cases_minimum ===")
    original_trace_counts = {}
    for i, log in enumerate(selected_logs):
        trace_count = log['case:concept:name'].nunique()
        original_trace_counts[selected_routines[i]] = trace_count
        print(f"  Routine {selected_routines[i]}: {trace_count} traces")
    
    # Find the minimum number of traces among all logs
    min_traces = min(original_trace_counts.values())
    print(f"Minimum traces across all routines: {min_traces}")
    
    next_case_id = 1
    final_trace_counts = {}
    
    for i, log in enumerate(selected_logs):
        # Select first min_traces case IDs
        unique_cases = log['case:concept:name'].unique()
        selected_case_ids = unique_cases[:min_traces]
        
        log = log[log['case:concept:name'].isin(selected_case_ids)].copy()
        
        # Create mapping to new sequential integer IDs
        case_mapping = {old: new for old, new in zip(selected_case_ids, range(next_case_id, next_case_id + len(selected_case_ids)))}
        log['case:concept:name'] = log['case:concept:name'].map(case_mapping)
        selected_logs[i] = log
        next_case_id += len(selected_case_ids)
        
        # Get the actual final trace count after filtering and renumbering
        final_trace_count = log['case:concept:name'].nunique()
        final_trace_counts[selected_routines[i]] = final_trace_count
        
        print(f"  DEBUG - Routine {selected_routines[i]}: Original={original_trace_counts[selected_routines[i]]}, Selected={len(selected_case_ids)}, Final={final_trace_count}")
    
    print(f"\n=== AFTER select_and_renumber_cases_minimum ===")
    for i, routine in enumerate(selected_routines):
        original_count = original_trace_counts[routine]
        final_count = final_trace_counts[routine]
        reduction = original_count - final_count
        reduction_percent = (reduction / original_count * 100) if original_count > 0 else 0
        print(f"  Routine {routine}: {original_count} → {final_count} traces (reduced by {reduction}, {reduction_percent:.1f}%)")
    
    print(f"\n=== SUMMARY ===")
    total_original = sum(original_trace_counts.values())
    total_final = sum(final_trace_counts.values())
    total_reduction = total_original - total_final
    print(f"  Total traces: {total_original} → {total_final} (reduced by {total_reduction})\n")
    
    return selected_logs


def adjust_timestamps_to_target(selected_logs, target_datetime_str):
    """Adjust timestamps to target datetime"""
    adjusted_logs = []
    target_datetime = pd.to_datetime(target_datetime_str)
    for log in selected_logs:
        timestamps = pd.to_datetime(log['time:timestamp'])
        time_deltas = timestamps - timestamps.min()
        log = log.copy()
        log['time:timestamp'] = target_datetime + time_deltas
        adjusted_logs.append(log)
    return adjusted_logs


def force_first_trace_to_target(adjusted_logs, target_datetime_str):
    """Force first trace to start at target datetime"""
    target_datetime = pd.to_datetime(target_datetime_str)
    new_logs = []
    for log in adjusted_logs:
        case_starts = log.groupby('case:concept:name')['time:timestamp'].min()
        first_event_time = case_starts.min()
        shift = target_datetime - first_event_time
        log = log.copy()
        log['time:timestamp'] = pd.to_datetime(log['time:timestamp']) + shift
        new_logs.append(log)
    return new_logs


def calculate_inter_trace_gaps(adjusted_logs):
    """Calculate gaps between traces"""
    inter_trace_gaps = {}
    for idx, log in enumerate(adjusted_logs):
        trace_times = []
        for case_id in log['case:concept:name'].unique():
            case_df = log[log['case:concept:name'] == case_id]
            start = case_df['time:timestamp'].min()
            end = case_df['time:timestamp'].max()
            trace_times.append((start, end))
        trace_times.sort()
        gaps = []
        for i in range(1, len(trace_times)):
            prev_end = trace_times[i-1][1]
            curr_start = trace_times[i][0]
            gap = (curr_start - prev_end).total_seconds()
            if gap > 0:
                gaps.append(gap)
        inter_trace_gaps[idx] = gaps
    variances = {idx: pd.Series(gaps).var() for idx, gaps in inter_trace_gaps.items()}
    return inter_trace_gaps, variances


def fit_best_distribution(data):
    """Fit best statistical distribution to data"""
    distributions = [stats.norm, stats.lognorm, stats.expon, stats.gamma, stats.beta]
    best_fit = None
    best_p = -1
    best_params = None
    for dist in distributions:
        try:
            params = dist.fit(data)
            D, p = stats.kstest(data, dist.name, args=params)
            if p > best_p:
                best_p = p
                best_fit = dist
                best_params = params
        except Exception:
            continue
    return best_fit, best_params


def fit_reference_distribution(time_diffs, variances, adjusted_logs, variance_criteria):
    """Fit reference distribution based on variance criteria"""
    reference_log_idx = max(variances, key=variances.get) if variance_criteria == 'max' else min(variances, key=variances.get)
    reference_log = adjusted_logs[reference_log_idx]
    time_diff_list = time_diffs[reference_log_idx]
    sorted_time_diffs = sorted(time_diff_list)
    best_dist, best_params = fit_best_distribution(sorted_time_diffs) if sorted_time_diffs else (None, None)
    return reference_log_idx, best_dist, best_params, sorted_time_diffs


def perform_first_round_shift(adjusted_logs, reference_log_idx, best_dist, best_params, sorted_time_diffs):
    """Perform first round of trace shifting"""
    updated_logs = []
    for log_idx, log in enumerate(adjusted_logs):
        if log_idx == reference_log_idx:
            updated_logs.append(log)
            continue
        new_times = []
        case_starts = log.groupby('case:concept:name')['time:timestamp'].min()
        case_ids = case_starts.sort_values().index.tolist()
        last_end_time = None
        if sorted_time_diffs:
            try:
                lower_bound = max(0.0, float(np.percentile(sorted_time_diffs, 1)))
            except Exception:
                lower_bound = 0.0
            try:
                upper_bound = float(np.percentile(sorted_time_diffs, 99))
            except Exception:
                upper_bound = float(max(sorted_time_diffs))
            if not np.isfinite(upper_bound) or upper_bound <= 0:
                upper_bound = float(max(sorted_time_diffs)) if max(sorted_time_diffs) > 0 else 60.0
            upper_bound = float(min(upper_bound, 30 * 24 * 3600))
        else:
            lower_bound = 0.0
            upper_bound = 60.0
        for idx, case_id in enumerate(case_ids):
            case_df = log[log['case:concept:name'] == case_id].copy()
            case_df = case_df.sort_values(by='time:timestamp')
            if idx == 0:
                new_times.append(case_df)
                last_end_time = case_df['time:timestamp'].max()
            else:
                sampled_gap_seconds = None
                for _ in range(100):
                    try:
                        if best_dist is not None and best_params is not None:
                            candidate = float(best_dist.rvs(*best_params))
                        else:
                            candidate = float(random.choice(sorted_time_diffs)) if sorted_time_diffs else 1.0
                    except Exception:
                        candidate = 1.0
                    if not np.isfinite(candidate):
                        continue
                    candidate = max(lower_bound, min(candidate, upper_bound))
                    if candidate >= 0:
                        sampled_gap_seconds = candidate
                        break
                if sampled_gap_seconds is None:
                    sampled_gap_seconds = min(upper_bound, 60.0)
                orig_start = case_df['time:timestamp'].min()
                if last_end_time is not None:
                    new_start = last_end_time + pd.Timedelta(seconds=float(sampled_gap_seconds))
                else:
                    new_start = case_df['time:timestamp'].min()
                shift = new_start - orig_start
                case_df['time:timestamp'] = case_df['time:timestamp'] + shift
                last_end_time = case_df['time:timestamp'].max()
                new_times.append(case_df)
        updated_log = pd.concat(new_times, ignore_index=True)
        updated_logs.append(updated_log)
    try:
        first_round = pd.concat(updated_logs, ignore_index=True)
        out_dir = os.path.join('out', 'logs')
        os.makedirs(out_dir, exist_ok=True)
        first_round.sort_values(by=['time:timestamp']).to_csv(os.path.join(out_dir, 'after_first_round.csv'), index=False)
    except Exception:
        pass
    return updated_logs


def perform_gap_scaling_to_target(updated_logs, reference_log_idx):
    """Perform gap scaling to target"""
    try:
        ref_log = updated_logs[reference_log_idx]
        T_target = ref_log['time:timestamp'].max()
    except Exception:
        return updated_logs
    final_logs = []
    for log_idx, log in enumerate(updated_logs):
        if log_idx == reference_log_idx:
            final_logs.append(log)
            continue
        case_starts = log.groupby('case:concept:name')['time:timestamp'].min().sort_values()
        ordered_case_ids = case_starts.index.tolist()
        if len(ordered_case_ids) <= 1:
            final_logs.append(log)
            continue
        trace_blocks = []
        for case_id in ordered_case_ids:
            case_df = log[log['case:concept:name'] == case_id].sort_values(by='time:timestamp').copy()
            trace_start = case_df['time:timestamp'].min()
            trace_end = case_df['time:timestamp'].max()
            trace_blocks.append((case_id, case_df, trace_start, trace_end))
        S_r = min(tb[2] for tb in trace_blocks)
        E_r = max(tb[3] for tb in trace_blocks)
        if pd.isna(S_r) or pd.isna(E_r) or T_target <= S_r:
            final_logs.append(log)
            continue
        gaps = []
        for i in range(1, len(trace_blocks)):
            prev_end = trace_blocks[i-1][3]
            next_start = trace_blocks[i][2]
            gap = (next_start - prev_end)
            gaps.append(gap)
        durations_sum = sum(((tb[3] - tb[2]) for tb in trace_blocks), pd.Timedelta(0))
        current_gaps_sum = sum(gaps, pd.Timedelta(0))
        desired_total_span = T_target - S_r
        new_gaps_sum = desired_total_span - durations_sum
        if new_gaps_sum < pd.Timedelta(0):
            new_gaps_sum = pd.Timedelta(0)
        if len(trace_blocks) == 1:
            final_logs.append(log)
            continue
        if current_gaps_sum > pd.Timedelta(0):
            scale = new_gaps_sum / current_gaps_sum if current_gaps_sum != pd.Timedelta(0) else 1.0
            new_gaps = [pd.to_timedelta(g * scale) for g in gaps]
        else:
            even_gap = new_gaps_sum / (len(trace_blocks) - 1) if (len(trace_blocks) - 1) > 0 else pd.Timedelta(0)
            new_gaps = [even_gap] * (len(trace_blocks) - 1)
        rebuilt = []
        first_case_df = trace_blocks[0][1].copy()
        rebuilt.append(first_case_df)
        rolling_end = first_case_df['time:timestamp'].max()
        for i in range(1, len(trace_blocks)):
            case_id, case_df, orig_start, orig_end = trace_blocks[i]
            desired_start = rolling_end + new_gaps[i-1]
            shift = desired_start - orig_start
            shifted_df = case_df.copy()
            shifted_df['time:timestamp'] = shifted_df['time:timestamp'] + shift
            rebuilt.append(shifted_df)
            rolling_end = shifted_df['time:timestamp'].max()
        final_logs.append(pd.concat(rebuilt, ignore_index=True))
    try:
        second_round = pd.concat(final_logs, ignore_index=True)
        out_dir = os.path.join('out', 'logs')
        os.makedirs(out_dir, exist_ok=True)
        second_round.sort_values(by=['time:timestamp']).to_csv(os.path.join(out_dir, 'after_second_round.csv'), index=False)
    except Exception:
        pass
    return final_logs


def shift_traces_by_distribution(adjusted_logs, reference_log_idx, best_dist, best_params, sorted_time_diffs):
    """Shift traces by distribution"""
    first_round_logs = perform_first_round_shift(adjusted_logs, reference_log_idx, best_dist, best_params, sorted_time_diffs)
    final_logs = perform_gap_scaling_to_target(first_round_logs, reference_log_idx)
    return final_logs


def compute_trace_interleaving_details(segmented_log):
    """
    Build per-trace overlap information.

    Returns a list of dictionaries with keys:
        - routine
        - case_id
        - start / end timestamps
        - overlap_routines (set of routines that overlap in time)
        - overlap_count (size of overlap_routines)
    """
    log_copy = segmented_log.copy()
    log_copy['time:timestamp'] = pd.to_datetime(log_copy['time:timestamp'])
    trace_intervals = []
    for routine in log_copy['routine_type'].unique():
        routine_log = log_copy[log_copy['routine_type'] == routine]
        for case_id in routine_log['case:concept:name'].unique():
            trace = routine_log[routine_log['case:concept:name'] == case_id]
            trace_intervals.append({
                'routine': routine,
                'case_id': case_id,
                'start': trace['time:timestamp'].min(),
                'end': trace['time:timestamp'].max()
            })
    details = []
    for idx, t1 in enumerate(trace_intervals):
        overlaps = set()
        for jdx, t2 in enumerate(trace_intervals):
            if idx == jdx:
                continue
            if t1['start'] <= t2['end'] and t2['start'] <= t1['end']:
                if t1['routine'] != t2['routine']:
                    overlaps.add(t2['routine'])
        details.append({
            'routine': t1['routine'],
            'case_id': t1['case_id'],
            'start': t1['start'],
            'end': t1['end'],
            'overlap_routines': overlaps,
            'overlap_count': len(overlaps)
        })
    return details


def count_trace_interleaving_cases(segmented_log):
    """Count trace interleaving cases"""
    segmented_log['time:timestamp'] = pd.to_datetime(segmented_log['time:timestamp'])
    trace_intervals = []
    for routine in segmented_log['routine_type'].unique():
        routine_log = segmented_log[segmented_log['routine_type'] == routine]
        for case_id in routine_log['case:concept:name'].unique():
            trace = routine_log[routine_log['case:concept:name'] == case_id]
            start = trace['time:timestamp'].min()
            end = trace['time:timestamp'].max()
            trace_intervals.append({
                'routine': routine,
                'case_id': case_id,
                'start': start,
                'end': end
            })
    interleaved_with = dict()
    for i, t1 in enumerate(trace_intervals):
        overlaps = set()
        for j, t2 in enumerate(trace_intervals):
            if i == j:
                continue
            if t1['start'] <= t2['end'] and t2['start'] <= t1['end']:
                if t1['routine'] != t2['routine']:
                    overlaps.add(t2['routine'])
        interleaved_with[(t1['routine'], t1['case_id'])] = overlaps
    count_0 = 0
    count_1 = 0
    count_2 = 0
    for overlap_set in interleaved_with.values():
        if len(overlap_set) == 0:
            count_0 += 1
        elif len(overlap_set) == 1:
            count_1 += 1
        elif len(overlap_set) >= 2:
            count_2 += 1
    print(f"\nNot interleaved: {count_0}")
    print(f"Interleaved with one: {count_1}")
    print(f"Interleaved with two or more: {count_2}\n")
    return {
        'not_interleaved': count_0,
        'interleaved_with_one': count_1,
        'interleaved_with_two_or_more': count_2
    }


def average_interleaving_per_log(segmented_log):
    """
    Compute the supervisor-requested metric: average # of interleaving routines per trace.

    Returns
    -------
    average : float
    std_dev : float
    per_trace_counts : List[int]
    """
    trace_details = compute_trace_interleaving_details(segmented_log)
    if not trace_details:
        return 0.0, 0.0, []
    counts = [detail['overlap_count'] for detail in trace_details]
    avg = float(np.mean(counts))
    std = float(np.std(counts, ddof=0))
    return avg, std, counts


def summarize_interleaving_samples(logs_by_class, doc_object=None, table_title="Average interleaving per class"):
    """
    logs_by_class: dict[label -> iterable of segmented logs].
        Each log item can be a pandas DataFrame or a CSV path.

    Returns (summary_df, samples_by_class, doc_object).
    """
    summary_rows = []
    samples_by_class = {}
    for class_name, log_items in logs_by_class.items():
        class_samples = []
        for log_item in log_items:
            if isinstance(log_item, pd.DataFrame):
                log_df = log_item
            elif isinstance(log_item, str):
                try:
                    log_df = pd.read_csv(log_item)
                except Exception as exc:
                    print(f"Skipping {log_item}: {exc}")
                    continue
            else:
                print(f"Unsupported log source for {class_name}: {type(log_item)}")
                continue
            avg, _, _ = average_interleaving_per_log(log_df)
            class_samples.append(avg)
        samples_by_class[class_name] = class_samples
        if class_samples:
            class_mean = float(np.mean(class_samples))
            class_std = float(np.std(class_samples, ddof=1)) if len(class_samples) > 1 else 0.0
        else:
            class_mean = class_std = 0.0
        summary_rows.append({
            'Class': class_name,
            '# Logs': len(class_samples),
            'Mean Avg Interleaving': class_mean,
            'Std Avg Interleaving': class_std
        })
    summary_df = pd.DataFrame(summary_rows)
    if not summary_df.empty:
        print(tabulate(summary_df, headers='keys', tablefmt='psql', showindex=False))
    if doc_object is not None and not summary_df.empty:
        doc_object.add_heading(table_title, level=2)
        table = doc_object.add_table(rows=1, cols=len(summary_df.columns))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for idx, col_name in enumerate(summary_df.columns):
            hdr[idx].text = str(col_name)
        for _, row in summary_df.iterrows():
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                if isinstance(value, float):
                    row_cells[idx].text = f"{value:.4f}"
                else:
                    row_cells[idx].text = str(value)
        doc_object.add_paragraph()
    return summary_df, samples_by_class, doc_object


def display_and_store_selected_routines(selected_routines, logs, document):
    """Display and store selected routines"""
    document.add_heading("Randomly Selected Logs", level=2)
    for num in selected_routines:
        text = f"Randomly select routine log {num} from log {logs[logs['routine_type'] == num]['log_number'].unique()[0]}"
        # display(HTML(styled_html))
        document.add_paragraph(text)
    return document


######################################################################################################################
####                                            Functions Called After process_random_logs                        ####
######################################################################################################################

def plot_interleaved_routines(final_log, numbers, doc_object, image_path="out/plots/interleaved_routines_plot.png", title="Interleaving of Selected Routines", duration='all', interleaved_only=False, clamp_to_threeway=False, plot_routine_level=False):
    """Plot interleaved routines"""
    if not os.path.exists(os.path.dirname(image_path)):
        os.makedirs(os.path.dirname(image_path))
        
    final_log['time:timestamp'] = pd.to_datetime(final_log['time:timestamp'])
    if duration in ['1h', '2h']:
        start_time = final_log['time:timestamp'].min()
        end_time = start_time + pd.Timedelta(hours=int(duration[0]))
        final_log = final_log[final_log['time:timestamp'] <= end_time]
    
    # Always plot individual traces
    trace_info, routines_to_plot = build_trace_intervals(final_log, numbers)
    
    if interleaved_only:
        interleaved_routines = set()
        interleaved_traces = []
        interleave_counts = defaultdict(int)
        for i in range(len(trace_info)):
            for j in range(i + 1, len(trace_info)):
                t1 = trace_info[i]
                t2 = trace_info[j]
                if t1['start'] <= t2['end'] and t2['start'] <= t1['end']:
                    interleaved_traces.extend([t1, t2])
                    interleaved_routines.update([t1['routine'], t2['routine']])
                    pair = tuple(sorted([t1['routine'], t2['routine']]))
                    interleave_counts[pair] += 1
        unique_traces = {(t['routine'], t['case_id']): t for t in interleaved_traces}
        plot_traces = list(unique_traces.values())
        routines_to_plot = sorted(interleaved_routines)
    else:
        plot_traces = trace_info
    
    plot_traces_gantt(plot_traces, routines_to_plot, image_path, f"{title} ({duration}) - {len(routines_to_plot)} routines", duration, cutoff_time=None)
    doc_object.add_heading(f"{title} ({duration})", level=2)
    doc_object.add_picture(image_path, width=Inches(6.5))
    doc_object.add_paragraph(f"📊 Total Routines Plotted: {len(routines_to_plot)}")
    if interleaved_only:
        doc_object.add_paragraph("🔁 Interleaving Counts Between Routines:")
        for (r1, r2), count in interleave_counts.items():
            doc_object.add_paragraph(f" - Routine {r1} & Routine {r2}: {count} overlapping traces")
    return doc_object


def build_trace_intervals(final_log, numbers):
    """Build trace intervals"""
    trace_info = []
    for routine in numbers:
        routine_log = final_log[final_log['routine_type'] == routine]
        for case_id in routine_log['case:concept:name'].unique():
            trace = routine_log[routine_log['case:concept:name'] == case_id]
            start = trace['time:timestamp'].min()
            end = trace['time:timestamp'].max()
            trace_info.append({
                'routine': routine,
                'case_id': case_id,
                'start': start,
                'end': end
            })
    routines_to_plot = sorted(set([t['routine'] for t in trace_info]))
    return trace_info, routines_to_plot


def plot_traces_gantt(trace_info, routines_to_plot, image_path, title, duration, cutoff_time=None):
    """Plot traces as Gantt chart"""
    plt.figure(figsize=(12, 5))
    routine_y_positions = {r: i for i, r in enumerate(routines_to_plot)}
    colors = plt.cm.get_cmap('tab10', len(routines_to_plot))
    for t in trace_info:
        y = routine_y_positions[t['routine']]
        line_start = t['start']
        line_end = t['end'] if cutoff_time is None else min(t['end'], cutoff_time)
        if cutoff_time is not None and line_start > cutoff_time:
            continue
        plt.hlines(
            y=y,
            xmin=line_start,
            xmax=line_end,
            colors=colors(routines_to_plot.index(t['routine'])),
            linewidth=6,
            label=f"Routine {t['routine']}"
        )
    handles, labels = plt.gca().get_legend_handles_labels()
    unique = dict(zip(labels, handles))
    plt.legend(unique.values(), unique.keys(), loc='upper center', bbox_to_anchor=(0.5, -0.15), ncol=len(unique))
    plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
    plt.gca().xaxis.set_major_locator(MaxNLocator(nbins=20))
    plt.yticks(list(routine_y_positions.values()), [f"Routine {r}" for r in routines_to_plot])
    plt.xlabel("Timestamp")
    plt.title(title)
    plt.grid(True, linestyle='--', alpha=0.5)
    plt.tight_layout()
    plt.savefig(image_path)
    plt.close()
    return image_path


def plot_trace_interleaving_cases(segmented_log, image_path="out/plots/trace_interleaving_cases_bar.png", title="Trace Interleaving Cases"):
    """Plot trace interleaving cases as bar chart"""
    if not os.path.exists(os.path.dirname(image_path)):
        os.makedirs(os.path.dirname(image_path))
    counts = count_trace_interleaving_cases(segmented_log)
    categories = list(counts.keys())
    values = list(counts.values())
    total = sum(values)
    plt.figure(figsize=(8, 5))
    bars = plt.bar(categories, values, color=['#4e79a7', '#f28e2b', '#e15759'])
    plt.title(f"{title} (Total: {total})")
    plt.ylabel("Number of Traces")
    plt.xlabel("Interleaving Type")
    plt.grid(axis='y', linestyle='--', alpha=0.5)
    for bar, value in zip(bars, values):
        percent = (value / total * 100) if total > 0 else 0
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, f'{int(value)} ({percent:.0f}%)',
                 ha='center', va='bottom', fontsize=12, fontweight='bold')
    plt.tight_layout()
    plt.savefig(image_path)
    plt.close()


def update_avg_counts(trace_interleaved_counts):
    # Compute totals
    total = (
        trace_interleaved_counts['not_interleaved']
        + trace_interleaved_counts['interleaved_with_one']
        + trace_interleaved_counts['interleaved_with_two_or_more']
    )

    # Add averages
    trace_interleaved_counts['avg_not_interleaved'] = (
        trace_interleaved_counts['not_interleaved'] / total
    )
    trace_interleaved_counts['avg_interleaved'] = (
        (trace_interleaved_counts['interleaved_with_one']
        + trace_interleaved_counts['interleaved_with_two_or_more']) / total
    )

    return trace_interleaved_counts


######################################################################################################################
####                                            Evaluation and Export Functions                                   ####
######################################################################################################################

def jaccard_coefficient(set1, set2):
    """Calculate Jaccard coefficient"""
    return len(set1 & set2) / len(set1 | set2) if set1 | set2 else 0


def extract_transition_labels(petri_net):
    """Extract transition labels from Petri net"""
    return set([t.label for t in petri_net.transitions if t.label])


def extract_activity_names_and_log_ids(activity_set):
    """Extract activity names and log IDs"""
    activity_names = []
    log_ids = set()
    
    for activity in activity_set:
        if '_' in activity:
            name, log_id = activity.rsplit('_', 1)
            activity_names.append(name)
            log_ids.add(log_id)
    
    return set(activity_names), log_ids


def evaluate_clusters_with_jc(cluster_dict, log_dataframe, logs_folder="logs", models_folder="GT_Models", display_results=True):
    """Evaluate clusters with Jaccard coefficient"""
    key_value = 1
    result_dict = {
        'Metrics':['Cluster ID', "# Activities", "Unique Activities", 'Trace Count', 'JC'],
    }
    for cluster_id, cluster_activities in cluster_dict.items():
        activity_names, routine_ids = extract_activity_names_and_log_ids(cluster_activities)
        log_numbers = [
            log_dataframe[log_dataframe['routine_type'] == int(routine_id)]['log_number'].unique()[0]
            for routine_id in routine_ids
        ]
        try:
            subset = log_dataframe[log_dataframe['routine_type'].isin([int(r) for r in routine_ids])]
            trace_count = int(subset['case:concept:name'].nunique())
        except Exception:
            trace_count = 0
        jc_scores = []
        jc_model_map = {}
        for log_number in log_numbers:
            model_path = os.path.join(models_folder, f"log{log_number}")
            if not os.path.exists(model_path):
                continue
            for file in os.listdir(model_path):
                if not file.endswith('.pnml'):
                    continue
                pnml_file_path = os.path.join(model_path, file)
                net, initial_marking, final_marking = pnml_importer.import_net(pnml_file_path)
                model_activities = extract_transition_labels(net)
                model_activities = set([act for act in list(model_activities) if not (act.startswith('start') or act.startswith('end'))])
                print(model_activities)
                jc = jaccard_coefficient(activity_names, model_activities)
                jc_scores.append(jc)
                jc_model_map[jc] = file
        max_jc = max(jc_scores) if jc_scores else 0
        best_jc_model = jc_model_map[max_jc] if max_jc in jc_model_map else "N/A"
        if display_results:
            print(f"Cluster {cluster_id}: Max JC = {max_jc:.4f}, Best GT Model: {best_jc_model}")
        result_dict[key_value] = [cluster_id, len(cluster_activities), len(activity_names), trace_count, max_jc]
        key_value += 1
    return result_dict


def evaluate_clusters_with_jc_variant(cluster_dict, log_dataframe, logs_folder="logs", models_folder="GT_Models", display_results=True):
    """Evaluate clusters with Jaccard coefficient, by computing JC for each trace with all GT models for its log number, retaining the best JC for each trace, and then averaging these best JCs for the cluster. Activity suffix is removed before JC."""
    key_value = 1
    result_dict = {
        'Metrics': ['Cluster ID', "# Activities", "Unique Activities", 'Trace Count', 'JC'],
    }
    for cluster_id, cluster_activities in cluster_dict.items():
        activity_names, routine_ids = extract_activity_names_and_log_ids(cluster_activities)
        routine_ids_int = [int(r) for r in routine_ids]
        # Prepare map: routine_type -> log_number for current cluster
        routine_log_map = {}
        for routine_id in routine_ids_int:
            matches = log_dataframe[log_dataframe['routine_type'] == routine_id]['log_number'].unique()
            if len(matches) > 0:
                routine_log_map[routine_id] = int(matches[0])
        # Filter dataframe once for all routines for this cluster
        subset = log_dataframe[log_dataframe['routine_type'].isin(routine_ids_int)]
        trace_count = int(subset['case:concept:name'].nunique())
        best_jcs = []
        # Group by (routine_type, case:concept:name): each trace
        for (routine_type, trace_id), trace_df in subset.groupby(['routine_type', 'case:concept:name']):
            if routine_type not in routine_log_map:
                continue
            log_number = routine_log_map[routine_type]
            model_path = os.path.join(models_folder, f"log{log_number}")
            if not os.path.exists(model_path):
                continue
            pnml_files = [file for file in os.listdir(model_path) if file.endswith('.pnml')]
            if not pnml_files:
                continue
            # Remove routine type suffix from activities in this trace
            cleaned_activities = set()
            for act in trace_df['concept:name'].unique():
                if '_' in act:
                    base, suffix = act.rsplit('_', 1)
                    try:
                        int(suffix)
                        cleaned_activities.add(base)
                        continue
                    except ValueError:
                        pass
                cleaned_activities.add(act)
            # Compute JC to all GT models, keep the best
            jc_list = []
            for file in pnml_files:
                pnml_file_path = os.path.join(model_path, file)
                net, initial_marking, final_marking = pnml_importer.import_net(pnml_file_path)
                model_activities = extract_transition_labels(net)
                jc = jaccard_coefficient(cleaned_activities, model_activities)
                jc_list.append(jc)
            best_jc = max(jc_list) if jc_list else 0
            best_jcs.append(best_jc)
        avg_jc = np.mean(best_jcs) if best_jcs else 0
        if display_results:
            print(f"Cluster {cluster_id}: Avg best JC of traces = {avg_jc:.4f} for {len(best_jcs)} traces")
        result_dict[key_value] = [cluster_id, len(cluster_activities), len(activity_names), trace_count, avg_jc]
        key_value += 1
    return result_dict


def append_averages_to_results(results_dict, selected_rotines):
    """Append averages to results"""
    selected_rotines = [str(routine) for routine in selected_rotines]
    routine_types_str = ', '.join(selected_rotines)
    simple_avg_label = f"Simple Average ({routine_types_str})"
    weighted_avg_label = f"Weighted Average ({routine_types_str})"
    headers = results_dict.get('Metrics', [])
    
    try:
        activities_idx = headers.index('# Activities')
    except ValueError:
        activities_idx = 1
    try:
        unique_activities_idx = headers.index('Unique Activities')
    except ValueError:
        unique_activities_idx = 2
    try:
        trace_count_idx = headers.index('Trace Count')
    except ValueError:
        trace_count_idx = 3
    try:
        jc_idx = headers.index('JC')
    except ValueError:
        jc_idx = len(headers) - 1

    simple_avg_row = [simple_avg_label]
    weighted_avg_row = [weighted_avg_label]
    total_activities = sum(results_dict[key][activities_idx] for key in results_dict if isinstance(key, int))
    total_unique_activities = sum(results_dict[key][unique_activities_idx] for key in results_dict if isinstance(key, int))
    total_trace_count = sum(results_dict[key][trace_count_idx] for key in results_dict if isinstance(key, int))
    simple_avg_row.append(str(total_activities))
    weighted_avg_row.append(str(total_activities))
    simple_avg_row.append(str(total_unique_activities))
    weighted_avg_row.append(str(total_unique_activities))
    simple_avg_row.append(str(total_trace_count))
    weighted_avg_row.append(str(total_trace_count))
    
    metric_values = [results_dict[key][jc_idx] for key in results_dict if isinstance(key, int)]
    simple_average = sum(metric_values) / len(metric_values) if metric_values else 0
    weighted_average = sum(
        results_dict[key][jc_idx] * results_dict[key][activities_idx]
        for key in results_dict if isinstance(key, int)
    ) / total_activities if total_activities else 0
    simple_avg_row.append(str(simple_average))
    weighted_avg_row.append(str(weighted_average))
    int_keys = [key for key in results_dict if isinstance(key, int)]
    next_index = (max(int_keys) + 1) if int_keys else 1
    results_dict[next_index] = simple_avg_row
    results_dict[next_index + 1] = weighted_avg_row
    return results_dict


def display_and_export_results_doc(results_dict, doc_object, table_title="Cluster Evaluation Summary"):
    """Display and export results to document"""
    result_df = pd.DataFrame(results_dict)
    result_df = result_df.set_index('Metrics')
    result_df = result_df.T
    print(tabulate(result_df, headers="keys", tablefmt="psql"))
    doc_object.add_heading(table_title, level=2)
    table = doc_object.add_table(rows=1, cols=len(result_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(result_df.columns):
        hdr_cells[i].text = str(col_name)
    for _, row in result_df.iterrows():
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = f"{cell_value:.4f}" if isinstance(cell_value, float) else str(cell_value)
    doc_object.add_paragraph()
    return doc_object


def collect_simple_average_rows_over_iterations(results_dicts):
    """Collect simple average rows over iterations"""
    rows = []
    for idx, results_dict in enumerate(results_dicts):
        for key, row in results_dict.items():
            if isinstance(row[0], str) and row[0].startswith("Simple Average"):
                rows.append([f"Iter {idx+1}"] + row)
                break
    if not rows:
        return []
    
    num_cols = len(rows[0])
    mean_row = ["Overall Average", "Overall Mean"]
    for col in range(2, num_cols):
        col_values = []
        for row in rows:
            try:
                col_values.append(float(row[col]))
            except Exception:
                col_values.append(0)
        mean_val = sum(col_values) / len(col_values) if col_values else 0
        if all(isinstance(v, int) or (isinstance(v, float) and v.is_integer()) for v in col_values):
            mean_val = int(round(mean_val))
        mean_row.append(str(mean_val))
    rows.append(mean_row)
    return rows


def collect_iteration_metadata(all_results, all_iteration_metadata):
    """Collect iteration summary metadata for the new sheet with separate rows for each routine"""
    iteration_summary_rows = []
    
    for i, (results_dict, metadata) in enumerate(zip(all_results, all_iteration_metadata)):
        if metadata is None:
            continue
            
        # Extract data from metadata
        selected_routines = metadata.get('selected_routines', [])
        traces_before = metadata.get('traces_before_renumbering', {})
        traces_after = metadata.get('traces_after_renumbering', {})
        interleaving_counts = metadata.get('interleaving_counts', {})
        
        # Create summary row
        num_routines = len(selected_routines)
        
        # Format interleaving information (same for all rows in this iteration)
        total_interleaving = sum(interleaving_counts.values()) if interleaving_counts else 0
        interleaving_str = f"Total: {total_interleaving}"
        if interleaving_counts:
            interleaving_str += f" (Not interleaved: {interleaving_counts.get('not_interleaved', 0)}, "
            interleaving_str += f"With 1: {interleaving_counts.get('interleaved_with_one', 0)}, "
            interleaving_str += f"With 2+: {interleaving_counts.get('interleaved_with_two_or_more', 0)})"
        
        # Create a row for each routine
        for routine in selected_routines:
            iteration_summary_rows.append([
                f"Iteration {i+1}",
                num_routines,
                str(routine),
                f"R{routine}: {traces_before.get(routine, 0)}",
                f"R{routine}: {traces_after.get(routine, 0)}",
                interleaving_str
            ])
        
        # Add empty row between iterations (except for the last iteration)
        if i < len(all_iteration_metadata) - 1:
            iteration_summary_rows.append([
                '', '', '', '', '', ''
            ])
    
    return iteration_summary_rows


def display_and_export_results_xlx(all_results, all_iteration_metadata=None, output_path="Transformed_Logs_and_Results/Our", filename="all_iterations_full_results.xlsx", include_metadata=True):
    """Display and export results to Excel with optional metadata sheet"""
    # Build all iteration results as a list of rows
    all_rows = []
    for i, results_dict in enumerate(all_results):
        df = pd.DataFrame(results_dict)
        df = df.set_index('Metrics').T.reset_index()
        # Remove both 'Iteration' and 'index' columns if they exist
        for col in ['Iteration', 'index']:
            if col in df.columns:
                df = df.drop(columns=[col])
        # Add a label row (first cell is label, rest are empty)
        label_row = [f"Iteration {i+1}"] + [''] * (df.shape[1] - 1)
        all_rows.append(label_row)
        # Add column names as a row for this iteration
        all_rows.append(list(df.columns))
        # Add the iteration's results
        all_rows.extend(df.values.tolist())
        # Add a blank row for separation
        all_rows.append([''] * df.shape[1])
    
    # Remove the last blank row if present
    if all_rows and all(all_rows[-1][j] == '' for j in range(len(all_rows[-1]))):
        all_rows = all_rows[:-1]
    
    # Create DataFrame for all iterations
    iteration_results_df = pd.DataFrame(all_rows)
    
    # Create summary DataFrame
    summary_rows = collect_simple_average_rows_over_iterations(all_results)
    summary_df = pd.DataFrame(
        summary_rows,
        columns=["Iteration", "Label", "# Activities", "Unique Activities", "Trace Count", "JC"]
    )
    
    # Create sheets to write
    sheets_to_write = {
        "All Iterations": iteration_results_df,
        "Summary": summary_df
    }
    
    # Add iteration metadata sheet if requested and metadata is provided
    if include_metadata and all_iteration_metadata is not None:
        iteration_metadata_rows = collect_iteration_metadata(all_results, all_iteration_metadata)
        metadata_df = pd.DataFrame(
            iteration_metadata_rows,
            columns=["Iteration", "Num Selected Routines", "Selected Routines", "Traces Before Renumbering", "Traces After Renumbering", "Interleaving Counts"]
        )
        sheets_to_write["Iteration Metadata"] = metadata_df
    
    # Write all sheets to Excel
    with pd.ExcelWriter(f"{output_path}/{filename}") as writer:
        for sheet_name, df in sheets_to_write.items():
            if sheet_name == "All Iterations":
                df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
            else:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    sheet_names = list(sheets_to_write.keys())
    print(f"Results exported to {filename} with sheets: {', '.join(sheet_names)}")


def run_multiple_iterations_with_metadata(logs, document, num_iterations=5, target_datetime_str="2024-06-07 21:34:25.493914+00:00", variance_criteria='max'):
    """Run multiple iterations and collect metadata for comprehensive Excel export"""
    all_results = []
    all_iteration_metadata = []
    
    for iteration in range(num_iterations):
        print(f"\n=== Starting Iteration {iteration + 1} ===")
        
        # Process logs for this iteration
        result = process_random_logs(
            logs=logs,
            document=document,
            target_datetime_str=target_datetime_str,
            variance_criteria=variance_criteria
        )
        
        segment_log, unsegment_log, selected_routines, document, iteration_metadata = result
        
        if segment_log is not None:
            # For now, create a dummy results_dict
            results_dict = {
                'Metrics': ['Cluster ID', "# Activities", "Unique Activities", 'Trace Count', 'JC'],
                1: [1, 10, 8, iteration_metadata['traces_after_renumbering'].get(selected_routines[0], 0), 0.85]
            }
            
            all_results.append(results_dict)
            all_iteration_metadata.append(iteration_metadata)
            
            print(f"Iteration {iteration + 1} completed successfully")
        else:
            print(f"Iteration {iteration + 1} failed")
            all_iteration_metadata.append(None)
    
    # Export all results with metadata
    display_and_export_results_xlx(
        all_results=all_results,
        all_iteration_metadata=all_iteration_metadata,
        output_path="Transformed_Logs_and_Results/Our",
        filename="all_iterations_with_metadata.xlsx"
    )
    
    return all_results, all_iteration_metadata


######################################################################################################################
####                                            Post-Export Analysis & Plots                                      ####
######################################################################################################################

def plot_jc_boxplot_by_non_interleaving_bins(results_path, bins=None, output_path="out/plots/jc_boxplot_by_non_interleaving_bins.png"):
    """Read results (xlsx multi-sheet or csv), compute non-interleaving % per iteration, bin, and plot JC boxplots.

    Inputs:
    - results_path: Path to exported results. Supports the xlsx generated by display_and_export_results_xlx.
                    For csv, expects an Iteration Metadata-like CSV with columns including 'Iteration' and
                    'Interleaving Counts'.
    - bins: Optional list of bin edges (percentages). Default is [0,10,20,30,40,50,60,70,80,90,100].
    - output_path: Where to save the resulting boxplot image.

    Behavior:
    - For xlsx: reads 'Summary' and 'Iteration Metadata' sheets.
    - For csv: reads a single CSV assumed to mirror the 'Iteration Metadata' sheet structure.
    - Computes per-iteration non-interleaving percentage and aligns with Simple Average JC for that iteration.
    - Bins iterations by non-interleaving percentage and plots boxplots of JC per bin.
    """
    if bins is None:
        bins = [0, 10, 20, 30, 40, 50, 60, 70, 80, 90, 100]

    # Ensure output directory exists
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    # Helper to parse interleaving counts string like:
    # "Total: 123 (Not interleaved: 45, With 1: 60, With 2+: 18)"
    def _parse_interleaving_counts(counts_str):
        if not isinstance(counts_str, str) or not counts_str:
            return None, None, None, None
        total_match = re.search(r"Total:\s*(\d+)", counts_str)
        not_inter_match = re.search(r"Not\s+interleaved:\s*(\d+)", counts_str)
        with1_match = re.search(r"With\s+1:\s*(\d+)", counts_str)
        with2_match = re.search(r"With\s*2\+?:\s*(\d+)", counts_str)
        try:
            total = int(total_match.group(1)) if total_match else None
            n0 = int(not_inter_match.group(1)) if not_inter_match else None
            n1 = int(with1_match.group(1)) if with1_match else None
            n2 = int(with2_match.group(1)) if with2_match else None
        except Exception:
            return None, None, None, None
        return total, n0, n1, n2

    def _iter_label_to_int(label):
        try:
            return int(str(label).split()[-1])
        except Exception:
            return None

    def _load_xlsx(results_path):
        summary_df = pd.read_excel(results_path, sheet_name="Summary")
        metadata_df = pd.read_excel(results_path, sheet_name="Iteration Metadata")
        # Per-iteration Simple Average JC
        summary_df = summary_df.copy()
        summary_df["Iteration"] = summary_df["Iteration"].astype(str)
        per_iter_summary = summary_df[summary_df["Iteration"].str.startswith("Iter ")]
        per_iter_simple_avg = per_iter_summary[per_iter_summary["Label"].astype(str).str.startswith("Simple Average")]
        if per_iter_simple_avg.empty:
            per_iter_simple_avg = per_iter_summary
        jc_by_iter = per_iter_simple_avg.assign(iter_index=per_iter_simple_avg["Iteration"].map(_iter_label_to_int))[["iter_index", "JC"]].dropna()
        # Interleaving % per iteration
        metadata_df = metadata_df.copy()
        metadata_df["Iteration"] = metadata_df["Iteration"].astype(str)
        metadata_df = metadata_df[metadata_df["Iteration"].str.startswith("Iteration ")]
        def first_counts(group):
            vals = group["Interleaving Counts"].dropna().astype(str)
            for v in vals:
                if v.strip():
                    return v
            return None
        counts_by_iter_label = metadata_df.groupby("Iteration").apply(first_counts).reset_index(name="counts")
        counts_by_iter = counts_by_iter_label.assign(iter_index=counts_by_iter_label["Iteration"].map(_iter_label_to_int))
        parsed = counts_by_iter["counts"].apply(_parse_interleaving_counts)
        counts_by_iter = counts_by_iter.join(pd.DataFrame(parsed.tolist(), columns=["total", "not_inter", "with1", "with2"]))
        counts_by_iter["non_interleaving_pct"] = counts_by_iter.apply(
            lambda r: (r["not_inter"] / (r["not_inter"] + r["with1"] + r["with2"]) * 100.0) if all(
                isinstance(r[k], (int, float)) and r[k] is not None for k in ["not_inter", "with1", "with2"]
            ) and (r["not_inter"] + r["with1"] + r["with2"]) > 0 else None,
            axis=1
        )
        return pd.merge(jc_by_iter, counts_by_iter[["iter_index", "non_interleaving_pct"]], on="iter_index", how="inner")

    def _load_csv(results_path):
        metadata_df = pd.read_csv(results_path)
        required_cols = {"Iteration", "Interleaving Counts"}
        if not required_cols.issubset(set(metadata_df.columns)):
            raise ValueError("CSV must contain at least 'Iteration' and 'Interleaving Counts' columns.")
        metadata_df = metadata_df.copy()
        metadata_df["Iteration"] = metadata_df["Iteration"].astype(str)
        metadata_df = metadata_df[metadata_df["Iteration"].str.startswith("Iteration ")]
        counts_by_iter_label = metadata_df.groupby("Iteration")["Interleaving Counts"].apply(lambda s: next((v for v in s.astype(str) if v.strip()), None)).reset_index(name="counts")
        counts_by_iter = counts_by_iter_label.assign(iter_index=counts_by_iter_label["Iteration"].map(_iter_label_to_int))
        parsed = counts_by_iter["counts"].apply(_parse_interleaving_counts)
        counts_by_iter = counts_by_iter.join(pd.DataFrame(parsed.tolist(), columns=["total", "not_inter", "with1", "with2"]))
        counts_by_iter["non_interleaving_pct"] = counts_by_iter.apply(
            lambda r: (r["not_inter"] / (r["not_inter"] + r["with1"] + r["with2"]) * 100.0) if all(
                isinstance(r[k], (int, float)) and r[k] is not None for k in ["not_inter", "with1", "with2"]
            ) and (r["not_inter"] + r["with1"] + r["with2"]) > 0 else None,
            axis=1
        )
        if "JC" in metadata_df.columns:
            jc_by_iter = metadata_df[metadata_df["Iteration"].str.startswith("Iteration ")].copy()
            jc_by_iter = jc_by_iter.assign(iter_index=jc_by_iter["Iteration"].map(_iter_label_to_int))
            jc_by_iter = jc_by_iter.groupby("iter_index")["JC"].mean().reset_index()
            return pd.merge(jc_by_iter, counts_by_iter[["iter_index", "non_interleaving_pct"]], on="iter_index", how="inner")
        else:
            raise ValueError("CSV mode requires either the xlsx 'Summary' sheet or a 'JC' column in the CSV.")

    # Load data depending on extension
    if results_path.lower().endswith(".xlsx"):
        df = _load_xlsx(results_path)
    elif results_path.lower().endswith(".csv"):
        df = _load_csv(results_path)
    else:
        raise ValueError("Unsupported file type. Provide a .xlsx (recommended) or a .csv with required columns.")

    # Drop rows with missing values
    df = df.dropna(subset=["JC", "non_interleaving_pct"]).copy()

    if df.empty:
        raise ValueError("No valid data to plot after merging JC and non-interleaving percentages.")

    # Bin by non-interleaving percentage
    # Labels like '10-20', '20-30', ...
    bin_labels = [f"{int(bins[i])}-{int(bins[i+1])}" for i in range(len(bins)-1)]
    df["bin"] = pd.cut(df["non_interleaving_pct"], bins=bins, labels=bin_labels, include_lowest=True, right=False)

    # Prepare data for boxplot: list of JC arrays per bin in the order of bin_labels
    # Ensure every bin appears on x-axis even if empty by inserting [np.nan]
    data_per_bin = []
    xticklabels = bin_labels[:]
    for label in bin_labels:
        subset = df[df["bin"] == label]["JC"].astype(float)
        if subset.empty:
            data_per_bin.append([np.nan])
        else:
            data_per_bin.append(subset.values)

    plt.figure(figsize=(12, 6))
    bp = plt.boxplot(data_per_bin, labels=xticklabels, showfliers=True)
    plt.xlabel("Non-interleaving % bins")
    plt.ylabel("Simple Average JC per iteration")
    plt.title("JC distribution by non-interleaving percentage bins")
    plt.grid(axis='y', linestyle='--', alpha=0.5)

    # Add bin counts above x-ticks
    # Count true (non-NaN) observations per bin
    counts = []
    for label in xticklabels:
        c = int(df[df["bin"] == label]["JC"].notna().sum())
        counts.append(c)
    ymin, ymax = plt.ylim()
    # Position text slightly above upper whisker line
    y_text = ymax - (ymax - ymin) * 0.02
    for idx, (label, c) in enumerate(zip(xticklabels, counts), start=1):
        plt.text(idx, y_text, f"n={c}", ha='center', va='bottom', fontsize=9)

    plt.tight_layout()
    plt.savefig(output_path)
    plt.close()
    print(f"Saved boxplot to {output_path}")



def read_simone_routines(routines_path):
    all_routines = []
    # Open the text file and read line by line
    with open(routines_path, "r") as file:
        for line in file:
            # Split the line into parts (splits by spaces by default)
            parts = line.strip().split()
            
            # Exclude the last item (supporting number) and store the routine list
            routine = parts[:-1]  # This excludes the last element
            all_routines.append(routine)  # Add the routine to the main list
    
    return all_routines


def read_dumas_routines(routines_path):
    all_routines = []
    # Open the text file and read line by line
    with open(routines_path, "r") as file:
        for line in file:
            # Use regex to extract the pattern between square brackets
            match = re.search(r"\[(.*?)\]", line)
            
            if match:
                # Split the pattern by commas to create a list of events
                routine = [event.strip() for event in match.group(1).split(',')]
                # Add the routine to the main list
                all_routines.append(routine)
    
    return all_routines


def calculate_average_jc_across_logs(results_base_path="out/results", technique="Our", noise_level="0.1", num_logs=100):
    """
    Calculate average JC across all logs by reading average rows from individual log result files.
    
    Parameters:
    - results_base_path: Base path where results are stored
    - noise_level: Noise level folder (e.g., "0.1")
    - num_logs: Number of logs to process (default 100)
    
    Returns:
    - Dictionary with statistics about JC values across all logs
    """
    jc_values = []
    log_stats = []
    
    print(f"Processing {num_logs} logs from {results_base_path}/Noise_{noise_level}/")
    
    for log_num in range(1, num_logs + 1):
        # Try both CSV and Excel formats
        csv_path = f"{results_base_path}/Noise_{noise_level}/{technique}_results_log{log_num}.csv"
        xlsx_path = f"{results_base_path}/Noise_{noise_level}/{technique}_results_log{log_num}.xlsx"
        
        file_path = None
        file_type = None
        
        # Check which file exists
        if os.path.exists(csv_path):
            file_path = csv_path
            file_type = 'csv'
        elif os.path.exists(xlsx_path):
            file_path = xlsx_path
            file_type = 'excel'
        else:
            print(f"Log {log_num}: No result file found (tried both .csv and .xlsx)")
            continue
        
        try:
            if file_type == 'csv':
                # Read CSV file
                df = pd.read_csv(file_path)
            else:
                # Read Excel file - try to find the summary sheet
                try:
                    # First try to read the 'Summary' sheet
                    df = pd.read_excel(file_path, sheet_name='Summary')
                except:
                    # If no Summary sheet, read the first sheet
                    df = pd.read_excel(file_path, sheet_name=0)
            
            print(f"Log {log_num}: Reading {file_type} file - {file_path}")
            print(f"Log {log_num}: Columns found: {list(df.columns)}")
            print(f"Log {log_num}: Shape: {df.shape}")
            
            # Look for JC column (case insensitive)
            jc_column = None
            for col in df.columns:
                if 'jc' in col.lower():
                    jc_column = col
                    break
            
            if jc_column is None:
                print(f"Log {log_num}: No 'JC' column found. Available columns: {list(df.columns)}")
                continue
            
            print(f"Log {log_num}: Using JC column: '{jc_column}'")
            
            # Strategy 1: Look for the LAST row that contains "Average" (this should be the final average)
            text_columns = df.select_dtypes(include=['object']).columns
            avg_row = None
            
            if len(text_columns) > 0:
                # Find all rows containing "Average"
                avg_mask = df[text_columns].astype(str).apply(
                    lambda x: x.str.contains('Average', case=False, na=False)
                ).any(axis=1)
                avg_rows = df[avg_mask]
                
                if not avg_rows.empty:
                    # Take the LAST average row (should be the final calculated average)
                    avg_row = avg_rows.iloc[-1]
                    print(f"Log {log_num}: Found {len(avg_rows)} average rows, using the LAST one")
                else:
                    # Strategy 2: Look for specific patterns in the last few rows
                    last_5_rows = df.tail(5)
                    for pattern in ['Simple Average', 'Weighted Average', 'Overall Average', 'Mean']:
                        pattern_mask = last_5_rows[text_columns].astype(str).apply(
                            lambda x: x.str.contains(pattern, case=False, na=False)
                        ).any(axis=1)
                        if pattern_mask.any():
                            avg_row = last_5_rows[pattern_mask].iloc[-1]  # Take the last matching row
                            print(f"Log {log_num}: Found pattern '{pattern}' in last 5 rows")
                            break
                    
                    # Strategy 3: If no pattern found, use the last row (assuming it's the average)
                    if avg_row is None:
                        avg_row = df.iloc[-1]
                        print(f"Log {log_num}: Using the last row as average")
            
            if avg_row is not None:
                # Get the single JC value from the average row
                jc_value = pd.to_numeric(avg_row[jc_column], errors='coerce')
                
                if not pd.isna(jc_value):
                    jc_values.append(jc_value)
                    log_stats.append({
                        'log_number': log_num,
                        'avg_jc': jc_value,
                        'file_type': file_type
                    })
                    print(f"Log {log_num}: Average JC = {jc_value:.4f} (single value from last average row)")
                else:
                    print(f"Log {log_num}: No valid JC value found in average row")
                    print(f"Log {log_num}: JC value in average row: {avg_row[jc_column]}")
            else:
                print(f"Log {log_num}: No average row found")
                # Debug: show last few rows
                print(f"Log {log_num}: Last 3 rows:")
                print(df.tail(3).to_string())
                
        except FileNotFoundError:
            print(f"Log {log_num}: File not found - {file_path}")
        except Exception as e:
            print(f"Log {log_num}: Error reading file - {str(e)}")
            import traceback
            traceback.print_exc()
    
    if not jc_values:
        print("No valid JC values found across any logs!")
        return None
    
    # Calculate overall statistics
    jc_array = np.array(jc_values)
    overall_stats = {
        'total_logs_processed': len(jc_values),
        'overall_mean_jc': np.mean(jc_array),
        'overall_std_jc': np.std(jc_array),
        'overall_min_jc': np.min(jc_array),
        'overall_max_jc': np.max(jc_array),
        'overall_median_jc': np.median(jc_array),
        'individual_log_stats': log_stats,
        'all_jc_values': jc_values
    }
    
    print(f"\n=== OVERALL STATISTICS ===")
    print(f"Total logs processed: {overall_stats['total_logs_processed']}")
    print(f"Overall Mean JC: {overall_stats['overall_mean_jc']:.4f}")
    print(f"Overall Std JC: {overall_stats['overall_std_jc']:.4f}")
    print(f"Overall Min JC: {overall_stats['overall_min_jc']:.4f}")
    print(f"Overall Max JC: {overall_stats['overall_max_jc']:.4f}")
    print(f"Overall Median JC: {overall_stats['overall_median_jc']:.4f}")
    
    return overall_stats


def export_overall_jc_summary(overall_stats, output_path="out/results", noise_level="0.1"):
    """
    Export overall JC summary to Excel file.
    
    Parameters:
    - overall_stats: Dictionary returned from calculate_average_jc_across_logs
    - output_path: Base output path
    - noise_level: Noise level for filename
    """
    if overall_stats is None:
        print("No statistics to export!")
        return
    
    # Create output directory
    os.makedirs(output_path, exist_ok=True)
    
    # Create summary DataFrame
    summary_data = {
        'Metric': [
            'Total Logs Processed',
            'Overall Mean JC',
            'Overall Std JC', 
            'Overall Min JC',
            'Overall Max JC',
            'Overall Median JC'
        ],
        'Value': [
            overall_stats['total_logs_processed'],
            f"{overall_stats['overall_mean_jc']:.4f}",
            f"{overall_stats['overall_std_jc']:.4f}",
            f"{overall_stats['overall_min_jc']:.4f}",
            f"{overall_stats['overall_max_jc']:.4f}",
            f"{overall_stats['overall_median_jc']:.4f}"
        ]
    }
    
    summary_df = pd.DataFrame(summary_data)
    
    # Create individual log statistics DataFrame
    individual_df = pd.DataFrame(overall_stats['individual_log_stats'])
    
    # Export to Excel with multiple sheets
    output_file = f"{output_path}/Overall_JC_Summary_Noise_{noise_level}.xlsx"
    
    with pd.ExcelWriter(output_file) as writer:
        summary_df.to_excel(writer, sheet_name='Overall Summary', index=False)
        individual_df.to_excel(writer, sheet_name='Individual Log Stats', index=False)
    
    print(f"Overall JC summary exported to: {output_file}")
    
    return output_file


def plot_jc_distribution_across_logs(overall_stats, output_path="out/plots", noise_level="0.1"):
    """
    Create plots showing JC distribution across all logs.
    
    Parameters:
    - overall_stats: Dictionary returned from calculate_average_jc_across_logs
    - output_path: Path to save plots
    - noise_level: Noise level for filename
    """
    if overall_stats is None:
        print("No statistics to plot!")
        return
    
    # Create output directory
    os.makedirs(output_path, exist_ok=True)
    
    jc_values = overall_stats['all_jc_values']
    
    # Create histogram
    plt.figure(figsize=(10, 6))
    plt.hist(jc_values, bins=20, alpha=0.7, color='skyblue', edgecolor='black')
    plt.axvline(overall_stats['overall_mean_jc'], color='red', linestyle='--', 
                label=f"Mean: {overall_stats['overall_mean_jc']:.4f}")
    plt.axvline(overall_stats['overall_median_jc'], color='green', linestyle='--', 
                label=f"Median: {overall_stats['overall_median_jc']:.4f}")
    plt.xlabel('Average JC per Log')
    plt.ylabel('Frequency')
    plt.title(f'Distribution of Average JC Across {overall_stats["total_logs_processed"]} Logs (Noise {noise_level})')
    plt.legend()
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    
    hist_path = f"{output_path}/jc_distribution_noise_{noise_level}.png"
    plt.savefig(hist_path)
    plt.close()
    
    # Create box plot
    plt.figure(figsize=(8, 6))
    plt.boxplot(jc_values, vert=True)
    plt.ylabel('Average JC per Log')
    plt.title(f'Box Plot of Average JC Across {overall_stats["total_logs_processed"]} Logs (Noise {noise_level})')
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    
    box_path = f"{output_path}/jc_boxplot_noise_{noise_level}.png"
    plt.savefig(box_path)
    plt.close()
    
    print(f"Plots saved to: {hist_path} and {box_path}")
    
    return hist_path, box_path


def analyze_jc_across_all_logs(results_base_path="out/results", technique="Our", noise_level="0.1", num_logs=100, 
                             export_results=True, create_plots=True):
    """
    Convenience function that combines all steps to analyze JC across all logs.
    
    This function:
    1. Reads average JC values from all individual log result files
    2. Calculates overall statistics
    3. Exports summary to Excel (optional)
    4. Creates distribution plots (optional)
    
    Parameters:
    - results_base_path: Base path where results are stored
    - noise_level: Noise level folder (e.g., "0.1")
    - num_logs: Number of logs to process (default 100)
    - export_results: Whether to export results to Excel (default True)
    - create_plots: Whether to create distribution plots (default True)
    
    Returns:
    - Dictionary with overall statistics
    """
    print(f"=== ANALYZING JC ACROSS ALL LOGS ===")
    print(f"Base path: {results_base_path}")
    print(f"Noise level: {noise_level}")
    print(f"Number of logs: {num_logs}")
    print("=" * 50)
    
    # Step 1: Calculate average JC across all logs
    overall_stats = calculate_average_jc_across_logs(
        results_base_path=results_base_path,
        technique=technique,
        noise_level=noise_level,
        num_logs=num_logs
    )
    
    if overall_stats is None:
        print("No valid data found. Exiting.")
        return None
    
    # Step 2: Export results to Excel (if requested)
    if export_results:
        print("\n=== EXPORTING RESULTS ===")
        excel_file = export_overall_jc_summary(
            overall_stats=overall_stats,
            output_path=results_base_path,
            noise_level=noise_level
        )
        print(f"Excel summary saved to: {excel_file}")
    
    # Step 3: Create plots (if requested)
    if create_plots:
        print("\n=== CREATING PLOTS ===")
        plot_paths = plot_jc_distribution_across_logs(
            overall_stats=overall_stats,
            output_path="out/plots",
            noise_level=noise_level
        )
        if plot_paths:
            print(f"Plots created: {plot_paths}")
    
    print("\n=== ANALYSIS COMPLETE ===")
    print(f"Successfully processed {overall_stats['total_logs_processed']} logs")
    print(f"Overall average JC: {overall_stats['overall_mean_jc']:.4f}")
    
    return overall_stats


def create_discovered_routines_csv(base_path="Transformed_Logs_and_Results/arebmann/Transformed_Log_With_Noise_0.1", 
                                 num_logs=100, prediction_value=-1):
    """
    Read log variants from input directories and create corresponding CSV files 
    in the Discovered_Routines directory structure.
    
    Parameters:
    - base_path: Base path containing the log directories
    - num_logs: Number of logs to process (default 100)
    - prediction_value: Always set to 1 as specified
    
    Creates CSV files with format:
    prediction,bound
    1,<number_of_rows_in_input_file>
    """
    import glob
    
    print(f"=== CREATING DISCOVERED ROUTINES CSV FILES ===")
    print(f"Base path: {base_path}")
    print(f"Number of logs: {num_logs}")
    print(f"Prediction value: {prediction_value}")
    print("=" * 60)
    
    successful_logs = 0
    failed_logs = []
    
    for log_num in range(1, num_logs + 1):
        # Input directory path
        input_dir = f"{base_path}/log{log_num}"
        
        # Output directory path
        output_dir = f"{base_path}/Discovered_Routines/log{log_num}"
        
        try:
            # Check if input directory exists
            if not os.path.exists(input_dir):
                print(f"Log {log_num}: Input directory not found - {input_dir}")
                failed_logs.append(f"Log {log_num}: Input directory not found")
                continue
            
            # Find all CSV files in the input directory
            input_files = glob.glob(os.path.join(input_dir, '*.csv'))
            
            if not input_files:
                print(f"Log {log_num}: No files found in input directory")
                failed_logs.append(f"Log {log_num}: No files found")
                continue
            
            # Create output directory
            os.makedirs(output_dir, exist_ok=True)
            
            # Process each file individually to keep all variants
            files_processed = 0
            for file_path in input_files:
                try:
                    df = pd.read_csv(file_path)
                    rows_in_file = len(df)
                    print(f"Log {log_num}: File {os.path.basename(file_path)} has {rows_in_file} rows")
                    
                    # Create output filename based on input filename
                    input_filename = os.path.basename(file_path)
                    # Add _pred before .csv extension
                    output_filename = input_filename.replace('.csv', '_pred.csv')
                    output_file = f"{output_dir}/{output_filename}"
                    
                    # Create DataFrame with the required format
                    csv_data = {
                        'prediction': [prediction_value],
                        'bound': [rows_in_file]
                    }
                    df_output = pd.DataFrame(csv_data)
                    
                    # Save to CSV
                    df_output.to_csv(output_file, index=False)
                    
                    print(f"Log {log_num}: Created {output_file} with prediction={prediction_value}, bound={rows_in_file}")
                    files_processed += 1
                    
                except Exception as e:
                    print(f"Log {log_num}: Error processing {os.path.basename(file_path)} - {str(e)}")
                    continue
            
            if files_processed > 0:
                successful_logs += 1
            else:
                print(f"Log {log_num}: No files were successfully processed")
                failed_logs.append(f"Log {log_num}: No files processed")
            
        except Exception as e:
            print(f"Log {log_num}: Error processing - {str(e)}")
            failed_logs.append(f"Log {log_num}: {str(e)}")
            continue
    
    # Summary
    print(f"\n=== SUMMARY ===")
    print(f"Successfully processed: {successful_logs} logs")
    print(f"Failed: {len(failed_logs)} logs")
    
    if failed_logs:
        print(f"\nFailed logs:")
        for failure in failed_logs:
            print(f"  - {failure}")
    
    return {
        'successful_logs': successful_logs,
        'failed_logs': len(failed_logs),
        'failed_details': failed_logs
    }


def create_discovered_routines_csv_advanced(base_path="Transformed_Logs_and_Results/arebmann/Transformed_Log_With_Noise_0.1", 
                                          num_logs=100, prediction_value=1, 
                                          output_filename="discovered_routines.csv"):
    """
    Advanced version with more options for creating discovered routines CSV files.
    Only processes CSV files.
    
    Parameters:
    - base_path: Base path containing the log directories
    - num_logs: Number of logs to process (default 100)
    - prediction_value: Always set to -1 as specified
    - output_filename: Name of the output CSV file (fallback only)
    """
    import glob
    
    print(f"=== CREATING DISCOVERED ROUTINES CSV FILES (ADVANCED) ===")
    print(f"Base path: {base_path}")
    print(f"Number of logs: {num_logs}")
    print(f"Prediction value: {prediction_value}")
    print(f"Output filename: {output_filename}")
    print("=" * 70)
    
    successful_logs = 0
    failed_logs = []
    detailed_stats = []
    
    for log_num in range(1, num_logs + 1):
        # Input directory path
        input_dir = f"{base_path}/log{log_num}"
        
        # Output directory path
        output_dir = f"{base_path}/Discovered_Routines/log{log_num}"
        
        log_stats = {
            'log_number': log_num,
            'input_dir': input_dir,
            'output_dir': output_dir,
            'files_found': 0,
            'total_rows': 0,
            'success': False,
            'error': None
        }
        
        try:
            # Check if input directory exists
            if not os.path.exists(input_dir):
                error_msg = f"Input directory not found - {input_dir}"
                print(f"Log {log_num}: {error_msg}")
                log_stats['error'] = error_msg
                failed_logs.append(f"Log {log_num}: {error_msg}")
                detailed_stats.append(log_stats)
                continue
            
            # Find all CSV files in the input directory
            input_files = glob.glob(os.path.join(input_dir, '*.csv'))
            
            log_stats['files_found'] = len(input_files)
            
            if not input_files:
                error_msg = "No CSV files found in input directory"
                print(f"Log {log_num}: {error_msg}")
                log_stats['error'] = error_msg
                failed_logs.append(f"Log {log_num}: {error_msg}")
                detailed_stats.append(log_stats)
                continue
            
            # Create output directory
            os.makedirs(output_dir, exist_ok=True)
            
            # Process each file individually to keep all variants
            files_processed = 0
            file_details = []
            total_rows = 0
            
            for file_path in input_files:
                try:
                    df = pd.read_csv(file_path)
                    rows_in_file = len(df)
                    total_rows += rows_in_file
                    
                    # Create output filename based on input filename
                    input_filename = os.path.basename(file_path)
                    # Add _pred before .csv extension
                    output_filename = input_filename.replace('.csv', '_pred.csv')
                    output_file = f"{output_dir}/{output_filename}"
                    
                    # Create DataFrame with the required format
                    csv_data = {
                        'prediction': [prediction_value],
                        'bound': [rows_in_file]
                    }
                    df_output = pd.DataFrame(csv_data)
                    
                    # Save to CSV
                    df_output.to_csv(output_file, index=False)
                    
                    print(f"Log {log_num}: Created {output_file}")
                    print(f"  - File: {os.path.basename(file_path)}")
                    print(f"  - Rows: {rows_in_file}")
                    print(f"  - Prediction: {prediction_value}, Bound: {rows_in_file}")
                    
                    file_details.append({
                        'filename': os.path.basename(file_path),
                        'output_file': output_filename,
                        'rows': rows_in_file
                    })
                    
                    files_processed += 1
                    
                except Exception as e:
                    print(f"Log {log_num}: Error processing {os.path.basename(file_path)} - {str(e)}")
                    continue
            
            log_stats['total_rows'] = total_rows
            log_stats['file_details'] = file_details
            log_stats['files_processed'] = files_processed
            
            if files_processed > 0:
                log_stats['success'] = True
                successful_logs += 1
                print(f"Log {log_num}: Successfully processed {files_processed} files")
            else:
                error_msg = "No files were successfully processed"
                print(f"Log {log_num}: {error_msg}")
                log_stats['error'] = error_msg
                failed_logs.append(f"Log {log_num}: {error_msg}")
            
        except Exception as e:
            error_msg = f"Error processing - {str(e)}"
            print(f"Log {log_num}: {error_msg}")
            log_stats['error'] = error_msg
            failed_logs.append(f"Log {log_num}: {error_msg}")
        
        detailed_stats.append(log_stats)
    
    # Summary
    print(f"\n=== DETAILED SUMMARY ===")
    print(f"Successfully processed: {successful_logs} logs")
    print(f"Failed: {len(failed_logs)} logs")
    
    if failed_logs:
        print(f"\nFailed logs:")
        for failure in failed_logs:
            print(f"  - {failure}")
    
    # Create summary report
    summary_df = pd.DataFrame(detailed_stats)
    summary_file = f"{base_path}/discovered_routines_summary.xlsx"
    
    try:
        with pd.ExcelWriter(summary_file) as writer:
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
        
        print(f"\nDetailed summary saved to: {summary_file}")
    except Exception as e:
        print(f"Warning: Could not save summary file - {str(e)}")
    
    return {
        'successful_logs': successful_logs,
        'failed_logs': len(failed_logs),
        'failed_details': failed_logs,
        'detailed_stats': detailed_stats,
        'summary_file': summary_file
    }